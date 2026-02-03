import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from io import BytesIO

# --- 1. 頁面配置 ---
st.set_page_config(page_title="馬尼 EP 戰略系統 v15.2", page_icon="🐴", layout="wide")

st.markdown("""
    <style>
    .main { background-color: #F8FAFC; color: #1E293B; }
    [data-testid="stSidebar"] { background-color: #FFFFFF !important; border-right: 1px solid #E2E8F0 !important; }
    .section-header { 
        font-size: 22px !important; color: #0F172A !important; font-weight: 800 !important; 
        margin-top: 30px !important; margin-bottom: 10px !important;
        display: flex; align-items: center; border-bottom: 2px solid #E2E8F0; padding-bottom: 8px;
    }
    .ep-tag {
        background-color: #D946EF; color: white; padding: 2px 8px; border-radius: 4px; 
        font-size: 12px; font-weight: bold; margin-left: 10px; vertical-align: middle;
    }
    .stButton>button { border-radius: 6px !important; font-weight: bold !important; }
    .ai-btn-small>div>button { 
        background-color: #0F172A !important; color: white !important; 
        font-size: 14px !important; height: 45px !important; width: 100%;
        border: 1px solid #0F172A;
    }
    .stTextArea textarea { font-family: 'Noto Sans TC', sans-serif; line-height: 1.6; }
    textarea::placeholder { color: #64748B !important; font-style: italic; font-size: 14px; }
    
    .date-badge {
        background-color: #DBEAFE; color: #1E40AF; padding: 4px 12px; 
        border-radius: 4px; font-weight: bold; font-size: 0.95em; margin-bottom: 8px; display: inline-block;
    }
    </style>
    """, unsafe_allow_html=True)

# --- 2. 馬尼 EP 邏輯庫 (白話實戰版) ---
STRATEGY_LOGIC = {
    "重點(節日)活動": {
        "p_purpose": "【不只看業績，要看導流】\n簡單說：這次活動除了要賺錢，重點是要讓多少人「截圖」跑來店裡？\n請寫下：1. 預計導流多少人進店？ 2. 預計成交幾單？",
        "p_core": "【別直接打折，要讓客人覺得賺到】\n這招叫「心理帳戶」。\n不要只寫「手機折一千」，要寫「買手機送價值 $1280 的大禮包」。\n你要用什麼「贈品」來包裝主商品，讓客人覺得不買是傻子？",
        "p_schedule": "（請點擊上方按鈕，系統會自動幫你算好 W1 到 W8 該做什麼）", 
        "p_sop": "【一句話決勝負】\n客人拿著手機截圖進來了，店員第一句話該講什麼？\n不要講「歡迎光臨」，要講一句能讓他「不好意思不買」或「立刻想試用」的話。",
    },
    "門市(快閃)活動": {
        "p_purpose": "【換現金，清空間】\n這波活動就是為了「活下去」和「清倉庫」。\n你要針對哪一批堆在倉庫長灰塵的貨進行處理？目標是變現多少現金？",
        "p_core": "【製造緊張感】\n這招叫「稀缺性」。\n不要讓活動看起來像沒人要的清倉。\n要說：「這批貨是廠商流出的，只有這週有，僅限學生/老客戶購買。」",
        "p_schedule": "（請點擊按鈕，生成 3+1 的快閃時程）",
        "p_sop": "【擺對位置就贏一半】\n這批貨要放在櫃檯最顯眼的地方。\n話術重點：「店長說這批賣完就不補了，你現在不拿等下可能就沒了。」",
    },
    "Apple發布銷售": {
        "p_purpose": "【天下武功，唯快不破】\nApple 發布後 72 小時是黃金期。\n你的目標是：要在別家店還在搞清楚規格時，你就已經把預購單拿在手上了。",
        "p_core": "【解決客人的選擇障礙】\n客人很焦慮，不知道該不該換。\n你要準備一張「懶人包圖表」，告訴他：「別想了，換這支就對了，因為...」",
        "p_schedule": "（以發表日為準，自動推算 T+24h 與 T+72h 關鍵節點）",
        "p_sop": "【專業權威感】\n店員要背熟規格差異。\n話術：「早買早享受，現在登記，你是第一批拿到的 VIP。」",
    }
}

MODULES = [
    ("p_inventory", "一、 庫存去化目標 (替死鬼名單)", "這次要犧牲哪支手機或配件來當「帶路雞」？(請列出型號/庫存量/成本，別心軟)"),
    ("p_purpose", "二、 活動目的與 KPI", "具體數字寫出來：導流人數？成交數？客單價要拉到多少？"),
    ("p_core", "三、 核心策略與誘餌", "你要用什麼好康(誘餌)把客人從手機螢幕前，拉到你的櫃檯前？"),
    ("p_schedule", "四、 作戰時程表", "請依照左側設定，點擊上方按鈕生成動態時程。"),
    ("p_sop", "五、 門市執行與話術", "客人進店第一句要說什麼？如何引導他一定要摸到商品？"),
    ("p_marketing", "六、 流量與素材策略", "社群文案要打什麼關鍵字？(讓客人搜尋得到的誘因)"),
    ("p_review", "七、 檢討與減法分析", "如果不幸失敗，每週花 15 分鐘檢查哪裡出錯？(點擊少？還是進店不買？)")
]

FIELDS = [m[0] for m in MODULES] + ["p_name", "p_proposer", "p_date", "p_type", "p_duration"]

# --- 3. 內建模範範本 (Demo Templates) ---
DEMO_TEMPLATES = {
    "🏆 示範：2026 母親節去化戰 (重點活動)": {
        "p_name": "2026 母親節 - 寵愛媽咪庫存清空戰",
        "p_proposer": "馬尼 EP",
        "p_type": "重點(節日)活動",
        "p_duration": 56,
        "p_inventory": "1. iPhone 15 Plus (粉色庫存過高)\n2. 上一季的按摩槍配件 (贈品用)",
        "p_purpose": "導流目標：300 人進店核銷。\n轉化率：20% (60單)。\nATV：$25,000。",
        "p_core": "【買大送小策略】\n購買指定機型，免費升級「媽咪放鬆大禮包」(其實是庫存配件)。\n標榜：讓媽媽換新機又桑一下，價值感 $1980。",
        "p_sop": "話術：「這組是母親節限定的，送完就沒了，你要不要先傳 Line 問一下媽媽喜歡粉色還是黃色？」",
        "p_marketing": "SEO：母親節禮物推薦、手機買一送一。\n社群：拍一段「媽媽收到爛禮物 vs 手機」的對比短片。",
        "p_review": "每週一早會檢視：廣告出去後，有沒有人截圖來問？沒有就改圖。",
        "p_schedule": "" # 讓使用者自己按按鈕生成
    },
    "⚡ 示範：月底配件快閃 (門市活動)": {
        "p_name": "月底救星 - 學生族快閃專案",
        "p_proposer": "馬尼 EP",
        "p_type": "門市(快閃)活動",
        "p_duration": 14,
        "p_inventory": "1. 舊款軍規防摔殼 (庫存 50 個)\n2. 傳輸線 (散裝)",
        "p_purpose": "目標：兩週內清掉 40 個殼。\n換取現金流：$20,000。",
        "p_core": "【身份稀缺性】\n憑「學生證」或「滿分考卷」，享銅板加購價。\n理由：慶祝開學季/期中考 (隨便找個理由)。",
        "p_sop": "陳列：放在櫃檯結帳區。\n話術：「同學，這款防摔殼原價 890，今天憑學生證只要 199，剩這幾個喔。」",
        "p_marketing": "IG 限動連發：倒數計時，每天拍貨架越來越空的樣子。",
        "p_review": "前三天賣不動，馬上改成「憑舊殼換購」。",
        "p_schedule": ""
    }
}

# 初始化 Session State
if 'templates_store' not in st.session_state:
    st.session_state.templates_store = DEMO_TEMPLATES.copy() # 載入示範檔

for f in FIELDS:
    if f not in st.session_state:
        if f == 'p_date': st.session_state[f] = datetime.now()
        elif f == 'p_duration': st.session_state[f] = 56
        else: st.session_state[f] = ""

# --- 4. 核心功能：動態時程計算引擎 (修復日期格式) ---
def calculate_dynamic_schedule(start_date, duration_days, mode):
    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt = start_dt + timedelta(days=duration_days)
    
    # 日期格式化 helper (只顯示 YYYY/MM/DD)
    def fmt(dt): return dt.strftime('%Y/%m/%d')
    def fmt_s(dt): return dt.strftime('%m/%d')
    
    schedule_text = ""
    
    if mode == "重點(節日)活動":
        p1_days = int(duration_days * 0.25)
        p2_days = int(duration_days * 0.25)
        p3_days = int(duration_days * 0.375)
        
        d1_end = start_dt + timedelta(days=p1_days)
        d2_start = d1_end + timedelta(days=1)
        d2_end = d2_start + timedelta(days=p2_days)
        d3_start = d2_end + timedelta(days=1)
        d3_end = d3_start + timedelta(days=p3_days)
        d4_start = d3_end + timedelta(days=1)
        
        schedule_text = (
            f"📅 活動總週期：{fmt(start_dt)} - {fmt(end_dt)} (共 {duration_days} 天)\n\n"
            f"🟢 第一階段：策略發想期 ({fmt_s(start_dt)} - {fmt_s(d1_end)})\n"
            f"   - 任務：PM 會議 I。決定要犧牲打擊的庫存品，定出KPI。\n\n"
            f"🟡 第二階段：企劃定案期 ({fmt_s(d2_start)} - {fmt_s(d2_end)})\n"
            f"   - 任務：素材製作完畢、SEO 文章上線、門市話術教學。\n\n"
            f"🔴 第三階段：執行曝光期 ({fmt_s(d3_start)} - {fmt_s(d3_end)})\n"
            f"   - 任務：廣告全開、門市強力推銷。每週檢討「點擊vs核銷」。\n\n"
            f"🔵 第四階段：收尾回收期 ({fmt_s(d4_start)} - {fmt(end_dt)})\n"
            f"   - 任務：Q4 減法分析。砍掉那些燒錢又沒用的動作。"
        )
        
    elif mode == "門市(快閃)活動":
        prep_days = 3
        exec_days = duration_days - prep_days
        d1_end = start_dt + timedelta(days=prep_days)
        d2_start = d1_end + timedelta(days=1)
        
        schedule_text = (
            f"📅 快閃週期：{fmt(start_dt)} - {fmt(end_dt)} (共 {duration_days} 天)\n\n"
            f"⚡ 第一階段：快速定案 ({fmt_s(start_dt)} - {fmt_s(d1_end)})\n"
            f"   - 任務：選好要清的貨，做一張圖，定一個讓店員好推的價格。\n\n"
            f"🔥 第二階段：精準投放與執行 ({fmt_s(d2_start)} - {fmt(end_dt)})\n"
            f"   - 任務：IG 限動狂發、貨架黃金位陳列。\n"
            f"   - 監控：前三天沒人買，立刻換話術或位置。"
        )

    elif mode == "Apple發布銷售":
        t_plus_1 = start_dt + timedelta(days=1)
        t_plus_3 = start_dt + timedelta(days=3)
        
        schedule_text = (
            f"📅 Apple 戰役啟動日：{fmt(start_dt)} (T-Day)\n\n"
            f"🧊 Pre-Event (準備期)：即日起至 {fmt_s(start_dt)}\n"
            f"   - 任務：先把「新舊機比較表」模板做好，等規格一出直接填空。\n\n"
            f"🚀 T+24h 爆發期 ({fmt_s(t_plus_1)})\n"
            f"   - 任務：懶人包上線、門市人員熟背規格差異。\n\n"
            f"💰 T+72h 轉化期 ({fmt_s(t_plus_3)})\n"
            f"   - 任務：收割預購單，確保第一批貨能滿足 VIP。"
        )
        
    return schedule_text

# --- 5. 側邊欄：戰略控制塔 ---
with st.sidebar:
    st.title("⚡ 馬尼 EP")
    st.caption("行銷活動規劃系統 v15.2")
    
    st.header("1. 作戰模式與週期")
    
    campaign_type = st.radio(
        "活動類型", 
        ["重點(節日)活動", "門市(快閃)活動", "Apple發布銷售"],
        key="p_type_selector"
    )
    
    # 預設天數邏輯
    default_days = 56
    if campaign_type == "門市(快閃)活動": default_days = 14
    elif campaign_type == "Apple發布銷售": default_days = 7
        
    duration = st.number_input("執行週期 (天)", value=default_days, min_value=1, step=1, key="p_duration_input")
    st.session_state.p_duration = duration
    
    st.info(f"目前設定：\n{campaign_type} | {duration} 天")

    st.divider()
    st.header("2. 載入戰略/示範")
    
    # 這裡會包含示範檔與使用者存檔
    tpl_options = ["(請選擇)"] + list(st.session_state.templates_store.keys())
    selected_tpl = st.selectbox("選擇範本", options=tpl_options)
    
    c1, c2 = st.columns(2)
    if c1.button("📥 讀取範本"):
        if selected_tpl != "(請選擇)":
            data = st.session_state.templates_store[selected_tpl]
            # 載入資料
            for k, v in data.items():
                if k in st.session_state: st.session_state[k] = v
            # 根據載入的範本類型，自動切換 Radio Button 顯示 (需透過 Session State workaround)
            # 但 Streamlit Radio 重整後會依據 key 狀態，這裡主要載入文字內容
            st.success(f"已載入：{selected_tpl}")
            st.rerun()
            
    if c2.button("💾 存檔"):
        if st.session_state.p_name:
            st.session_state.templates_store[f"{st.session_state.p_name}"] = {f: st.session_state[f] for f in FIELDS}
            st.success("已存檔")

# --- 6. 主工作區 ---
st.title(f"🚀 戰略規劃：{campaign_type}")
st.caption("「馬尼 EP 精神：講人話、做實事、看現金流。」")

col1, col2, col3 = st.columns([2, 1, 1])
with col1: st.text_input("專案名稱", key="p_name", placeholder="ex: 2026 母親節 - 庫存去化戰")
with col2: st.text_input("負責人 (PM)", key="p_proposer")
with col3: st.date_input("活動起始日", key="p_date")

st.divider()

# 取得邏輯指引
current_logic = STRATEGY_LOGIC.get(campaign_type, STRATEGY_LOGIC["重點(節日)活動"])

# 計算時程字串
dynamic_schedule_content = calculate_dynamic_schedule(st.session_state.p_date, st.session_state.p_duration, campaign_type)
date_str_display = st.session_state.p_date.strftime('%Y/%m/%d') # UI 顯示用

# 渲染模組
for fid, title, default_guide in MODULES:
    st.markdown(f'<div class="section-header">{title} <span class="ep-tag">{campaign_type}</span></div>', unsafe_allow_html=True)
    pulse_guide = current_logic.get(fid, default_guide)
    
    # 特殊處理 p_schedule
    if fid == "p_schedule":
        st.markdown(f'<div class="date-badge">📅 系統運算：依據 {date_str_display} 起跑，共 {st.session_state.p_duration} 天</div>', unsafe_allow_html=True)
        if st.button("⚡ 自動帶入系統計算的時程與任務", key="btn_auto_schedule"):
            st.session_state[fid] = dynamic_schedule_content
            st.rerun()
            
    c_input, c_ai = st.columns([3, 1])
    
    with c_input:
        st.text_area(
            label="hidden", 
            key=fid, 
            height=200 if fid == "p_schedule" else 150, 
            placeholder=f"【馬尼 EP 指導】\n{pulse_guide}", 
            label_visibility="collapsed"
        )

    with c_ai:
        st.markdown('<div class="ai-btn-small">', unsafe_allow_html=True)
        if st.button(f"⚡ EP 診斷", key=f"btn_{fid}"):
            st.session_state[fid] = (
                f"【馬尼 EP 診斷】\n"
                f"1. 這太文言文了，能不能講人話？\n"
                f"2. 這樣做真的能換到現金嗎？還是只是在燒工讀生薪水？\n"
                f"3. 具體一點，下一步要做什麼？\n"
                f"---\n(原內容)\n{st.session_state[fid]}"
            )
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
        
        with st.expander("戰略提示"):
            st.markdown(f"**核心邏輯：**\n{pulse_guide}")

# --- 7. 輸出報告 ---
def generate_docx():
    doc = Document()
    doc.add_heading(f'馬尼 EP 戰略報告 - {campaign_type}', 0)
    doc.add_paragraph(f"專案：{st.session_state.p_name} | PM：{st.session_state.p_proposer}")
    
    # 這裡也要修復日期顯示
    d_start = st.session_state.p_date.strftime('%Y/%m/%d')
    doc.add_paragraph(f"週期：{d_start} 起，共 {st.session_state.p_duration} 天")
    
    for fid, title, _ in MODULES:
        doc.add_heading(title, level=2)
        content = st.session_state[fid] if st.session_state[fid] else "（未填寫）"
        doc.add_paragraph(content)
        
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

st.divider()
if st.button("📄 產出馬尼 EP 戰略書 (.docx)", type="primary"):
    file_data = generate_docx()
    st.download_button(
        label="📥 下載檔案",
        data=file_data,
        file_name=f"MoneyEP_{st.session_state.p_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
